# Standard Libraries
import logging
import os
import random
import re

# Django Imports
from django.conf import settings

# 3rd Party Libraries
import docx
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.image.exceptions import UnrecognizedImageError
from docx.oxml.shared import OxmlElement, qn
from docx.shared import Inches, Pt
from docx.shared import RGBColor as DocxRgbColor
from docx.document import Document as DocumentObject
from lxml import etree

# Ghostwriter Libraries
from ghostwriter.modules.reportwriter.base import ReportExportTemplateError
from ghostwriter.modules.reportwriter.extensions import (
    IMAGE_EXTENSIONS,
    TEXT_EXTENSIONS,
)
from ghostwriter.modules.reportwriter.richtext.ooxml import (
    BaseHtmlToOOXML,
    parse_styles,
)

logger = logging.getLogger(__name__)


class HtmlToDocx(BaseHtmlToOOXML):
    """
    Converts HTML to a word document
    """
    doc: DocumentObject
    p_style: str

    def __init__(self, doc: DocumentObject, p_style: str):
        super().__init__()
        self.doc = doc
        self.p_style = p_style

    def text(self, el, *, par=None, style={}, **kwargs):
        # Process hyperlinks on top of the usual text rules
        if par is not None and style.get("hyperlink_url"):
            # For Word, this code is modified from this issue:
            #   https://github.com/python-openxml/python-docx/issues/384
            # Get an ID from the ``document.xml.rels`` file
            part = par.part
            r_id = part.relate_to(
                style["hyperlink_url"],
                docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK,
                is_external=True,
            )
            # Create the ``w:hyperlink`` tag and add needed values
            hyperlink = docx.oxml.shared.OxmlElement("w:hyperlink")
            hyperlink.set(
                docx.oxml.shared.qn("r:id"),
                r_id,
            )
            # Create the ``w:r`` and ``w:rPr`` elements
            new_run = docx.oxml.shared.OxmlElement("w:r")
            rPr = docx.oxml.shared.OxmlElement("w:rPr")
            new_run.append(rPr)
            self.text_tracking.append_text_to_run(new_run, str(el))
            hyperlink.append(new_run)
            # Create a new Run object and add the hyperlink into it
            run = par.add_run()
            run._r.append(hyperlink)
            # A workaround for the lack of a hyperlink style
            if "Hyperlink" in self.doc.styles:
                try:
                    run.style = "Hyperlink"
                except KeyError:
                    pass
            else:
                run.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
                run.font.underline = True
        else:
            super().text(el, par=par, style=style, **kwargs)

    def style_run(self, run, style):
        super().style_run(run, style)
        if style.get("inline_code"):
            try:
                run.style = "CodeInline"
            except KeyError:
                pass
            run.font.no_proof = True
        if style.get("highlight"):
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        if style.get("strikethrough"):
            run.font.strike = True
        if style.get("subscript"):
            run.font.subscript = True
        if style.get("superscript"):
            run.font.superscript = True
        if style.get("font_size"):
            run.font.size = Pt(style["font_size"])
        if "font_color" in style:
            run.font.color.rgb = DocxRgbColor(*style["font_color"])

    def tag_br(self, el, *, par=None, **kwargs):
        self.text_tracking.new_block()
        if "data-gw-pagebreak" in el.attrs:
            self.doc.add_page_break()
        elif par is not None:
            run = par.add_run()
            run.add_break()

    def _tag_h(self, el, **kwargs):
        heading_num = int(el.name[1:])
        self.text_tracking.new_block()
        heading_paragraph = self.doc.add_heading(el.text, heading_num)

        bookmark_name = el.attrs.get("data-bookmark", el.attrs.get("id"))
        if bookmark_name and heading_paragraph.runs:
            tag = heading_paragraph.runs[0]._r
            start = docx.oxml.shared.OxmlElement("w:bookmarkStart")
            start.set(docx.oxml.ns.qn("w:id"), str(self.current_bookmark_id))
            start.set(docx.oxml.ns.qn("w:name"), "_Ref" + bookmark_name)
            tag.insert(0, start)

            tag = heading_paragraph.runs[-1]._r
            end = docx.oxml.shared.OxmlElement("w:bookmarkEnd")
            end.set(docx.oxml.ns.qn("w:id"), str(self.current_bookmark_id))
            end.set(docx.oxml.ns.qn("w:name"), "_Ref" + bookmark_name)
            tag.append(end)
            self.current_bookmark_id += 1

    tag_h1 = _tag_h
    tag_h2 = _tag_h
    tag_h3 = _tag_h
    tag_h4 = _tag_h
    tag_h5 = _tag_h
    tag_h6 = _tag_h

    def tag_p(self, el, *, par=None, **kwargs):
        self.text_tracking.new_block()
        if par is not None:
            # <p> nested in another block element like blockquote, use or copy the paragraph object
            if any(run.text for run in par.runs):
                # Paragraph has things in it already, make a new one but copy the style
                par = self.doc.add_paragraph(style=par.style)
        else:
            # Top level <p>
            par = self.doc.add_paragraph()
            try:
                par.style = self.p_style
            except KeyError:
                pass

        par_classes = set(el.attrs.get("class", []))
        if "left" in par_classes:
            par.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if "center" in par_classes:
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if "right" in par_classes:
            par.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if "justify" in par_classes:
            par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        self.process_children(el, par=par, **kwargs)

    def tag_pre(self, el, *, par=None, **kwargs):
        if par is not None and not any(run.text for run in par.runs):
            # Use provided empty paragraph
            pass
        elif par is not None:
            # Nested in li or some other element, copy style
            par = self.doc.add_paragraph(style=par.style)
        else:
            par = self.doc.add_paragraph()

        par.alignment = WD_ALIGN_PARAGRAPH.LEFT

        if len(el.contents) == 1 and el.contents[0].name == "code":
            content_el = el.contents[0]
        else:
            content_el = el

        self.text_tracking.new_block()
        self.text_tracking.in_pre = True
        try:
            self.process_children(content_el, par=par, **kwargs)
        finally:
            self.text_tracking.in_pre = False

        try:
            par.style = "CodeBlock"
        except KeyError:
            for run in par.runs:
                run.font.name = "Courier New"
                run.font.no_proof = True

    def tag_ul(self, el, *, par=None, list_level=None, list_tracking=None, **kwargs):
        if list_tracking is None:
            list_tracking = ListTracking()
            assert list_level is None
            this_list_level = 0
        else:
            assert list_level is not None
            this_list_level = list_level + 1

        is_ordered = el.name == "ol"

        # Create paragraphs for each list item
        for child in el.children:
            if child.name != "li":
                # TODO: log
                continue
            par = self.doc.add_paragraph()
            self.text_tracking.new_block()
            list_tracking.add_paragraph(par, this_list_level, is_ordered)
            self.process_children(
                child.children,
                par=par,
                list_level=this_list_level,
                list_tracking=list_tracking,
                **kwargs,
            )

        if this_list_level == 0:
            list_tracking.create(self.doc)

    tag_ol = tag_ul

    def tag_blockquote(self, el, par=None, **kwargs):
        # TODO: if done in a list, this won't preserve the level.
        # Not sure how to do that, since this requires a new paragraph.
        par = self.doc.add_paragraph()
        self.text_tracking.new_block()
        try:
            par.style = "Blockquote"
        except KeyError:
            pass
        self.process_children(el.children, par=par, **kwargs)

    def tag_div(self, el, **kwargs):
        if "page-break" in el.attrs.get("class", []):
            self.text_tracking.new_block()
            self.doc.add_page_break()
        else:
            super().tag_div(el, **kwargs)

    def create_table(self, rows, cols, **kwargs):
        table = self.doc.add_table(rows=rows, cols=cols, style="Table Grid")
        table.autofit = True
        table.allow_autofit = True
        table._tblPr.xpath("./w:tblW")[0].attrib[
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type"
        ] = "auto"
        for row_idx, _ in enumerate(table.rows):
            for cell_idx, _ in enumerate(table.rows[row_idx].cells):
                table.rows[row_idx].cells[cell_idx]._tc.tcPr.tcW.type = "auto"
                table.rows[row_idx].cells[cell_idx]._tc.tcPr.tcW.w = 0
        return table

    def paragraph_for_table_cell(self, cell, td_el):
        def handle_style(key, value):
            if key == "background-color":
                shade = OxmlElement("w:shd")
                shade.set(qn("w:fill"), value.replace("#", ""))
                cell._tc.get_or_add_tcPr().append(shade)

        parse_styles(td_el.attrs.get("style", ""), handle_style)

        return next(iter(cell.paragraphs))


class HtmlToDocxWithEvidence(HtmlToDocx):
    """
    Augments `HtmlToDocx`, replacing marked spans with evidence figures, captions,
    and references.
    """

    def __init__(
        self,
        doc,
        *,
        p_style,
        evidence_image_width,
        evidences,
        figure_label: str,
        figure_prefix: str,
        figure_caption_location: str,
        table_label: str,
        table_prefix: str,
        table_caption_location: str,
        title_case_captions: bool,
        title_case_exceptions: list[str],
        border_color_width: tuple[str, float] | None,
    ):
        super().__init__(doc, p_style)
        self.evidences = evidences
        self.figure_label = figure_label
        self.figure_prefix = figure_prefix
        self.figure_caption_location = figure_caption_location
        self.table_label = table_label
        self.table_prefix = table_prefix
        self.table_caption_location = table_caption_location
        self.title_case_captions = title_case_captions
        self.title_case_exceptions = title_case_exceptions
        self.border_color_width = border_color_width
        self.evidence_image_width = evidence_image_width
        self.plural_acronym_pattern = re.compile(r"^[^a-z]+(:?s|'s)$")
        self.current_bookmark_id = 1000 # Hopefully won't conflict with templates

    def text(self, el, *, par=None, **kwargs):
        if par is not None and getattr(par, "_gw_is_caption", False):
            el = self.title_except(el)
        return super().text(el, par=par, **kwargs)

    def tag_span(self, el, *, par, **kwargs):
        if "data-gw-evidence" in el.attrs:
            try:
                evidence = self.evidences[int(el.attrs["data-gw-evidence"])]
            except (KeyError, ValueError):
                return
            self.make_evidence(par, evidence)
        elif "data-gw-caption" in el.attrs:
            ref_name = el.attrs["data-gw-caption"]
            self.make_caption(par, self.figure_label, ref_name or None)
            par.add_run(self.figure_prefix)
        elif "data-gw-ref" in el.attrs:
            ref_name = el.attrs["data-gw-ref"]
            self.text_tracking.force_emit_pending_segment_break()
            self.make_cross_ref(par, ref_name)
        else:
            super().tag_span(el, par=par, **kwargs)

    def tag_table(self, el, **kwargs):
        caption_el = kwargs.get("caption_el") or el.find("caption")
        caption_bookmark = kwargs.get("caption_bookmark")
        if self.table_caption_location == "top":
            self._mk_table_caption(caption_el, caption_bookmark)
        super().tag_table(el, **kwargs)
        if self.table_caption_location == "bottom":
            self._mk_table_caption(caption_el, caption_bookmark)

    def tag_div(self, el, **kwargs):
        if "richtext-evidence" in el.attrs.get("class", []):
            try:
                evidence = self.evidences[int(el.attrs["data-evidence-id"])]
            except (KeyError, ValueError):
                logger.exception("Could not get evidence")
                return

            par = self.doc.add_paragraph()
            self.make_evidence(par, evidence)
        else:
            super().tag_div(el, **kwargs)

    def _mk_table_caption(self, caption_el, caption_bookmark=None):
        par_caption = self.doc.add_paragraph()
        self.make_caption(par_caption, self.table_label, caption_bookmark, styles=["Quote", "Caption"])
        if caption_el is not None:
            par_caption.add_run(self.table_prefix)
            par_caption.add_run(self.title_except(caption_el.get_text()))

    def is_plural_acronym(self, word):
        """
        Check if a word is an all caps acronym that ends with "s" or "'s".
        """
        return re.match(self.plural_acronym_pattern, word)

    def title_except(self, s):
        """
        Title case the given string except for articles and words in the provided exceptions list and words in all caps.

        Ref: https://stackoverflow.com/a/3729957
        """
        if self.title_case_captions:
            word_list = re.split(" ", s)  # re.split behaves as expected
            final = []
            for word in word_list:
                final.append(
                    word
                    if word in self.title_case_exceptions or word.isupper() or self.is_plural_acronym(word)
                    else word.capitalize()
                )
            s = " ".join(final)
        return s

    def make_caption(self, par, label: str, ref: str | None = None, styles = ["Caption"]):
        par._gw_is_caption = True
        for style in styles:
            try:
                par.style = style
                break
            except KeyError:
                continue

        if ref:
            ref = f"_Ref{ref}"
        else:
            ref = f"_Ref{random.randint(10000000, 99999999)}"

        # Start a bookmark run with the figure label
        p = par._p
        bookmark_start = OxmlElement("w:bookmarkStart")
        bookmark_start.set(qn("w:name"), ref.replace(" ", "_"))
        bookmark_start.set(qn("w:id"), str(self.current_bookmark_id))
        p.append(bookmark_start)

        # Add the figure label
        par.add_run(label)

        # Append XML for a new field character run
        run = par.add_run()
        r = run._r
        fldChar = OxmlElement("w:fldChar")
        fldChar.set(qn("w:fldCharType"), "begin")
        r.append(fldChar)

        # Add field code instructions with ``instrText``
        run = par.add_run()
        r = run._r
        instrText = OxmlElement("w:instrText")
        # Sequential figure with arabic numbers
        instrText.text = f" SEQ {label} \\* ARABIC"
        r.append(instrText)

        # An optional ``separate`` value to enforce a space between label and number
        run = par.add_run()
        r = run._r
        fldChar = OxmlElement("w:fldChar")
        fldChar.set(qn("w:fldCharType"), "separate")
        r.append(fldChar)

        # Include ``#`` as a placeholder for the number when Word updates fields
        run = par.add_run("#")
        r = run._r
        # Close the field character run
        fldChar = OxmlElement("w:fldChar")
        fldChar.set(qn("w:fldCharType"), "end")
        r.append(fldChar)

        # End the bookmark after the number
        p = par._p
        bookmark_end = OxmlElement("w:bookmarkEnd")
        bookmark_end.set(qn("w:id"), str(self.current_bookmark_id))
        p.append(bookmark_end)

        self.current_bookmark_id += 1

    def make_evidence(self, par, evidence):
        file_path = settings.MEDIA_ROOT + "/" + evidence["path"]
        if not os.path.exists(file_path):
            raise FileNotFoundError(file_path)

        extension = file_path.split(".")[-1].lower()
        if extension in TEXT_EXTENSIONS:
            try:
                with open(file_path, "r", encoding="utf-8") as evidence_file:
                    evidence_text = evidence_file.read()
            except UnicodeDecodeError as err:
                logger.exception(
                    "Evidence file known as %s (%s) was not recognized as a %s file.",
                    evidence["friendly_name"],
                    file_path,
                    extension,
                )
                error_msg = (
                    f'The evidence file, `{evidence["friendly_name"]},` was not recognized as a UTF-8 encoded {extension} file. '
                    "Try opening it, exporting as desired type, and re-uploading it."
                )
                raise ReportExportTemplateError(error_msg) from err

            if self.figure_caption_location == "top":
                self._mk_figure_caption(par, evidence["friendly_name"], evidence["caption"])
                par = self.doc.add_paragraph()

            par.text = evidence_text
            par.alignment = WD_ALIGN_PARAGRAPH.LEFT
            try:
                par.style = "CodeBlock"
            except KeyError:
                pass

            if self.figure_caption_location == "bottom":
                par_caption = self.doc.add_paragraph()
                self._mk_figure_caption(par_caption, evidence["friendly_name"], evidence["caption"])

        elif extension in IMAGE_EXTENSIONS:
            if self.figure_caption_location == "top":
                self._mk_figure_caption(par, evidence["friendly_name"], evidence["caption"])
                par = self.doc.add_paragraph()

            par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = par.add_run()
            try:
                run.add_picture(file_path, width=Inches(self.evidence_image_width))
            except UnrecognizedImageError as e:
                logger.exception(
                    "Evidence file known as %s (%s) was not recognized as a %s file.",
                    evidence["friendly_name"],
                    file_path,
                    extension,
                )
                error_msg = (
                    f'The evidence file, `{evidence["friendly_name"]},` was not recognized as a {extension} file. '
                    "Try opening it, exporting as desired type, and re-uploading it."
                )
                raise ReportExportTemplateError(error_msg) from e

            if self.border_color_width is not None:
                border_color, border_width = self.border_color_width
                # Add the border – see Ghostwriter Wiki for documentation
                inline_class = run._r.xpath("//wp:inline")[-1]
                inline_class.attrib["distT"] = "0"
                inline_class.attrib["distB"] = "0"
                inline_class.attrib["distL"] = "0"
                inline_class.attrib["distR"] = "0"

                # Set the shape's "effect extent" attributes to the border weight
                effect_extent = OxmlElement("wp:effectExtent")
                effect_extent.set("l", str(border_width))
                effect_extent.set("t", str(border_width))
                effect_extent.set("r", str(border_width))
                effect_extent.set("b", str(border_width))
                # Insert just below ``<wp:extent>`` or it will not work
                inline_class.insert(1, effect_extent)

                # Find inline shape properties – ``pic:spPr``
                pic_data = run._r.xpath("//pic:spPr")[-1]
                # Assemble OXML for a solid border
                ln_xml = OxmlElement("a:ln")
                ln_xml.set("w", str(border_width))
                solidfill_xml = OxmlElement("a:solidFill")
                color_xml = OxmlElement("a:srgbClr")
                color_xml.set("val", border_color)
                solidfill_xml.append(color_xml)
                ln_xml.append(solidfill_xml)
                pic_data.append(ln_xml)

            if self.figure_caption_location == "bottom":
                par_caption = self.doc.add_paragraph()
                self._mk_figure_caption(par_caption, evidence["friendly_name"], evidence["caption"])

    def _mk_figure_caption(self, par_caption, ref: str | None, caption_text: str):
        self.make_caption(par_caption, self.figure_label, ref)
        par_caption.add_run(self.figure_prefix)
        par_caption.add_run(self.title_except(caption_text))

    def make_cross_ref(self, par, ref: str):
        # Start the field character run for the label and number
        run = par.add_run()
        r = run._r
        fldChar = OxmlElement("w:fldChar")
        fldChar.set(qn("w:fldCharType"), "begin")
        r.append(fldChar)

        # Add field code instructions with ``instrText`` that points to the target bookmark
        run = par.add_run()
        r = run._r
        instrText = OxmlElement("w:instrText")
        instrText.text = ' REF "_Ref{}" \\h '.format(ref.replace("\\", "\\\\").replace('"', '\\"').replace(" ", "_"))
        r.append(instrText)

        # An optional ``separate`` value to enforce a space between label and number
        run = par.add_run()
        r = run._r
        fldChar = OxmlElement("w:fldChar")
        fldChar.set(qn("w:fldCharType"), "separate")
        r.append(fldChar)

        # Add runs for the figure label and number
        run = par.add_run(self.figure_label)
        # This ``#`` is a placeholder Word will replace with the figure's number
        run = par.add_run("#")

        # Close the  field character run
        run = par.add_run()
        r = run._r
        fldChar = OxmlElement("w:fldChar")
        fldChar.set(qn("w:fldCharType"), "end")
        r.append(fldChar)


class ListTracking:
    """
    Tracks a list being created, and creates the abstract and concrete ooxml numbering for it
    """

    def __init__(self):
        self.paragraphs = []
        self.level_list_is_ordered = []

    @staticmethod
    def q_w(tag):
        return etree.QName("http://schemas.openxmlformats.org/wordprocessingml/2006/main", tag)

    def add_paragraph(self, pg, level: int, is_ordered: bool):
        """
        Adds a paragraph to the list, tracking the list level and ordered/unordered type
        """
        if level > len(self.level_list_is_ordered):
            raise Exception(
                "Tried to add level {} to a list with {} existing levels".format(level, len(self.level_list_is_ordered))
            )
        if level == len(self.level_list_is_ordered):
            self.level_list_is_ordered.append(is_ordered)
        self.paragraphs.append((pg, level, is_ordered))

    def create(self, doc):
        """
        Creates the numbering, if needed, and assigns it to each of the paragraphs registered by `add_paragraph`.
        """
        level_list_is_ordered = self.level_list_is_ordered

        # Create a new numbering
        try:
            numbering = doc.part.numbering_part.numbering_definitions._numbering
        except NotImplementedError as e:
            raise ReportExportTemplateError("Tried to use a list in a template without list styles") from e
        last_used_id = max(
            (int(id) for id in numbering.xpath("w:abstractNum/@w:abstractNumId")),
            default=-1,
        )
        abstract_numbering_id = last_used_id + 1

        abstract_numbering = numbering.makeelement(self.q_w("abstractNum"))
        abstract_numbering.set(self.q_w("abstractNumId"), str(abstract_numbering_id))

        multi_level_type = abstract_numbering.makeelement(self.q_w("multiLevelType"))
        multi_level_type.set(self.q_w("val"), "hybridMultilevel")
        abstract_numbering.append(multi_level_type)

        for level_num, is_ordered in enumerate(level_list_is_ordered):
            # TODO: vary bullets or numbers based on level
            level = abstract_numbering.makeelement(self.q_w("lvl"))
            level.set(self.q_w("ilvl"), str(level_num))

            start = level.makeelement(self.q_w("start"))
            start.set(self.q_w("val"), "1")
            level.append(start)

            num_fmt = level.makeelement(self.q_w("numFmt"))
            lvl_text = level.makeelement(self.q_w("lvlText"))
            if is_ordered:
                num_fmt.set(self.q_w("val"), "decimal")
                lvl_text.set(self.q_w("val"), "%{}.".format(level_num + 1))
            else:
                num_fmt.set(self.q_w("val"), "bullet")
                lvl_text.set(self.q_w("val"), "")
                # lvl_text.set(self.q_w("val"), "X")
            level.append(num_fmt)
            level.append(lvl_text)

            prp = level.makeelement(self.q_w("pPr"))
            ind = prp.makeelement(self.q_w("ind"))
            ind.set(self.q_w("left"), str((level_num + 1) * 720))
            ind.set(self.q_w("hanging"), "360")
            prp.append(ind)
            level.append(prp)

            if not is_ordered:
                rpr = level.makeelement(self.q_w("rPr"))
                fonts = rpr.makeelement(self.q_w("rFonts"))
                fonts.set(self.q_w("ascii"), "Symbol")
                fonts.set(self.q_w("hAnsi"), "Symbol")
                fonts.set(self.q_w("hint"), "default")
                rpr.append(fonts)
                level.append(rpr)

            abstract_numbering.append(level)

        numbering.insert(0, abstract_numbering)
        numbering_id = numbering.add_num(abstract_numbering_id).numId

        for par, level, is_ordered in self.paragraphs:
            try:
                par.style = "Number List" if is_ordered else "Bullet List"
            except KeyError:
                try:
                    par.style = "ListParagraph"
                except KeyError:
                    pass
            par._p.get_or_add_pPr().get_or_add_numPr().get_or_add_numId().val = numbering_id
            par._p.get_or_add_pPr().get_or_add_numPr().get_or_add_ilvl().val = level
