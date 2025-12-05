"""Tests for footnote functionality using local python-docx fork."""

import os
import shutil
import tempfile

from django.test import TestCase

from docx import Document

from ghostwriter.modules.reportwriter.richtext.docx import HtmlToDocx


class FootnoteCreationTests(TestCase):
    """Test footnote creation with python-docx."""

    def setUp(self):
        """Set up test fixtures."""
        self.temp_dir = tempfile.mkdtemp()

    def tearDown(self):
        """Clean up temporary files."""
        shutil.rmtree(self.temp_dir, ignore_errors=True)

    def _get_baseline_footnote_count(self):
        """Get the number of footnotes in a blank document (separator footnotes)."""
        doc = Document()
        output_path = os.path.join(self.temp_dir, "baseline.docx")
        doc.save(output_path)
        reopened = Document(output_path)
        return len(reopened.footnotes)

    def test_create_document_with_footnote(self):
        """Test creating a Word document with a footnote."""
        document = Document()

        # Add a paragraph
        paragraph = document.add_paragraph("This is a paragraph with a footnote")

        # Add a footnote to the paragraph
        new_footnote = paragraph.add_footnote()
        new_footnote.add_paragraph("This is the footnote text.")

        # Verify footnote was created
        self.assertIsNotNone(new_footnote)
        self.assertEqual(len(new_footnote.paragraphs), 1)

        # Save to test output for inspection
        output_path = os.path.join(self.temp_dir, "test_footnote.docx")
        document.save(output_path)

        # Verify the file was created
        self.assertTrue(os.path.exists(output_path))

        # Reopen and verify footnote exists
        reopened_doc = Document(output_path)
        self.assertIsNotNone(reopened_doc.footnotes)
        self.assertGreater(len(reopened_doc.footnotes), 0)

    def test_create_multiple_footnotes(self):
        """Test creating multiple footnotes in a document."""
        baseline_count = self._get_baseline_footnote_count()

        document = Document()

        # Add first paragraph with footnote
        para1 = document.add_paragraph("First paragraph")
        footnote1 = para1.add_footnote()
        footnote1.add_paragraph("First footnote text.")

        # Add second paragraph with footnote
        para2 = document.add_paragraph("Second paragraph")
        footnote2 = para2.add_footnote()
        footnote2.add_paragraph("Second footnote text.")

        # Verify both footnotes were created
        self.assertIsNotNone(footnote1)
        self.assertIsNotNone(footnote2)

        # Footnote IDs should be different
        self.assertNotEqual(footnote1.id, footnote2.id)

        # Save to test output for inspection
        output_path = os.path.join(self.temp_dir, "test_multiple_footnotes.docx")
        document.save(output_path)
        self.assertTrue(os.path.exists(output_path))

        # Reopen and verify both footnotes exist (plus baseline separator footnotes)
        reopened_doc = Document(output_path)
        self.assertEqual(len(reopened_doc.footnotes), baseline_count + 2)

    def test_footnote_with_multiple_paragraphs(self):
        """Test creating a footnote with multiple paragraphs."""
        document = Document()

        paragraph = document.add_paragraph("Paragraph with multi-paragraph footnote")

        # Add footnote with multiple paragraphs
        footnote = paragraph.add_footnote()
        footnote.add_paragraph("First paragraph in footnote.")
        footnote.add_paragraph("Second paragraph in footnote.")

        # Verify footnote has multiple paragraphs
        self.assertEqual(len(footnote.paragraphs), 2)

        # Save to test output for inspection
        output_path = os.path.join(self.temp_dir, "test_multi_para_footnote.docx")
        document.save(output_path)
        self.assertTrue(os.path.exists(output_path))

    def test_access_footnotes_from_paragraph(self):
        """Test accessing footnotes from a paragraph."""
        document = Document()

        # Create a paragraph with a footnote
        paragraph = document.add_paragraph("Text with footnote reference")
        new_footnote = paragraph.add_footnote()
        new_footnote.add_paragraph("The footnote content.")

        # Save to test output for inspection
        output_path = os.path.join(self.temp_dir, "test_access_footnotes.docx")
        document.save(output_path)

        reopened_doc = Document(output_path)
        reopened_para = reopened_doc.paragraphs[0]

        # Access footnotes from paragraph
        footnotes = reopened_para.footnotes
        self.assertIsNotNone(footnotes)
        self.assertGreater(len(footnotes), 0)

        # Verify footnote properties
        footnote = footnotes[0]
        self.assertIsNotNone(footnote.id)
        self.assertGreater(len(footnote.paragraphs), 0)

    def test_access_all_document_footnotes(self):
        """Test accessing all footnotes from document."""
        baseline_count = self._get_baseline_footnote_count()

        document = Document()

        # Create multiple paragraphs with footnotes
        num_footnotes = 3
        for i in range(num_footnotes):
            para = document.add_paragraph(f"Paragraph {i + 1}")
            footnote = para.add_footnote()
            footnote.add_paragraph(f"Footnote {i + 1} content.")

        # Save to test output for inspection
        output_path = os.path.join(self.temp_dir, "test_all_footnotes.docx")
        document.save(output_path)

        reopened_doc = Document(output_path)

        # Access all footnotes (includes baseline separator footnotes)
        all_footnotes = reopened_doc.footnotes
        self.assertEqual(len(all_footnotes), baseline_count + num_footnotes)

        # Verify each footnote has an ID and paragraphs
        for footnote in all_footnotes:
            self.assertIsNotNone(footnote.id)
            self.assertGreaterEqual(len(footnote.paragraphs), 0)

    def test_footnote_ids_are_sequential(self):
        """Test that footnote IDs are assigned sequentially."""
        document = Document()

        footnote_ids = []
        for i in range(3):
            para = document.add_paragraph(f"Paragraph {i + 1}")
            footnote = para.add_footnote()
            footnote.add_paragraph(f"Footnote {i + 1}")
            footnote_ids.append(footnote.id)

        # Save to test output for inspection
        output_path = os.path.join(self.temp_dir, "test_sequential_footnotes.docx")
        document.save(output_path)

        # Verify IDs are unique and sequential
        self.assertEqual(len(footnote_ids), len(set(footnote_ids)))
        for i in range(1, len(footnote_ids)):
            self.assertGreater(footnote_ids[i], footnote_ids[i - 1])


class FootnoteRichTextConversionTests(TestCase):
    """Test footnote conversion from HTML rich text to DOCX."""

    def setUp(self):
        """Set up test fixtures."""
        self.temp_dir = tempfile.mkdtemp()

    def tearDown(self):
        """Clean up temporary files."""
        shutil.rmtree(self.temp_dir, ignore_errors=True)

    def test_html_to_docx_with_footnotes(self):
        """
        Test converting HTML with footnotes to DOCX using HtmlToDocx.
        """
        # Create a new document
        doc = Document()

        # Sample HTML content with footnotes (simulating rich text from the editor)
        html_content = """
        <p>This is an executive summary with important findings.<span class="footnote">Source: Internal Security Assessment, 2024</span></p>

        <p>During the assessment, several critical vulnerabilities were identified.<span class="footnote">See Appendix A for full vulnerability details.</span> The team discovered issues in the authentication system<span class="footnote">Authentication bypass via SQL injection in login form.</span> and the session management module.</p>

        <h2>Key Findings</h2>

        <p>The following critical issues require immediate attention:</p>

        <ul>
            <li>SQL Injection vulnerability in user login<span class="footnote">CVE-2024-1234 - CVSS Score 9.8</span></li>
            <li>Cross-Site Scripting (XSS) in comment fields</li>
            <li>Insecure Direct Object References<span class="footnote">Allows unauthorized access to other users' data.</span></li>
        </ul>

        <p>We recommend prioritizing remediation based on risk severity.<span class="footnote">Risk ratings based on CVSS v3.1 scoring methodology.</span></p>

        <h2>Recommendations</h2>

        <p>Implement input validation and parameterized queries<span class="footnote">OWASP recommends using prepared statements for all database queries.</span> to prevent injection attacks. Additionally, deploy a Web Application Firewall (WAF)<span class="footnote">Consider solutions like ModSecurity or cloud-based WAF services.</span> for defense in depth.</p>
        """

        # Run the HTML to DOCX converter
        HtmlToDocx.run(html_content, doc, None)

        # Save to temp file and reopen to verify footnotes
        output_path = os.path.join(self.temp_dir, "test_footnotes.docx")
        doc.save(output_path)

        # Reopen and verify footnotes exist
        reopened_doc = Document(output_path)
        self.assertIsNotNone(reopened_doc.footnotes)

        # Count actual footnotes (excluding separator footnotes which have id <= 0)
        actual_footnotes = [fn for fn in reopened_doc.footnotes if fn.id > 0]
        self.assertEqual(len(actual_footnotes), 8, "Expected 8 footnotes in the document")

        # Verify footnote content (strip leading space from footnoteRef marker)
        footnote_texts = [
            fn.paragraphs[0].text.strip() if fn.paragraphs else ""
            for fn in actual_footnotes
        ]
        self.assertIn("Source: Internal Security Assessment, 2024", footnote_texts)
        self.assertIn("CVE-2024-1234 - CVSS Score 9.8", footnote_texts)
        self.assertIn(
            "OWASP recommends using prepared statements for all database queries.",
            footnote_texts,
        )

    def test_html_to_docx_footnotes_in_table(self):
        """Test footnotes within table cells."""
        doc = Document()

        # HTML with a table containing footnotes
        html_content = """
        <p>Summary of Findings:</p>

        <table>
            <thead>
                <tr>
                    <th>Finding</th>
                    <th>Severity</th>
                    <th>Status</th>
                </tr>
            </thead>
            <tbody>
                <tr>
                    <td>SQL Injection<span class="footnote">Found in login and search forms.</span></td>
                    <td>Critical</td>
                    <td>Open</td>
                </tr>
                <tr>
                    <td>XSS Vulnerability<span class="footnote">Reflected XSS in query parameters.</span></td>
                    <td>High</td>
                    <td>In Progress</td>
                </tr>
                <tr>
                    <td>Missing CSRF Token<span class="footnote">State-changing operations lack protection.</span></td>
                    <td>Medium</td>
                    <td>Fixed</td>
                </tr>
            </tbody>
        </table>

        <p>See the detailed findings section for remediation steps.<span class="footnote">Remediation guidance follows OWASP best practices.</span></p>
        """

        HtmlToDocx.run(html_content, doc, None)

        # Save to temp file and reopen to verify footnotes
        output_path = os.path.join(self.temp_dir, "test_table_footnotes.docx")
        doc.save(output_path)

        reopened_doc = Document(output_path)
        actual_footnotes = [fn for fn in reopened_doc.footnotes if fn.id > 0]
        self.assertEqual(len(actual_footnotes), 4, "Expected 4 footnotes in the document")
